# Import necessary libraries
import requests
import json
import time
import numpy as np
import os
import docx
import aiohttp
import matplotlib.pyplot as plt
import scipy.stats as stats
from scipy.stats import norm, expon
from statistics import mean, stdev
from io import BytesIO
import concurrent.futures
import asyncio
import aiofiles

class QueryPerformanceAnalyzer:
    def __init__(self, query_file_path):
        self.query_file_path = query_file_path
        self.gateway_url = "http://localhost:33000/api/gql"
        self.headers = {"Content-Type": "application/json"}

    async def runSingleTestSet(self, query, num_queries):
        times = []
        error = False

        for _ in range(num_queries):
            async with aiohttp.ClientSession() as session:
                start_time = time.time()
                async with session.post(self.gateway_url, json=query, headers=self.headers) as response:
                    end_time = time.time()
                    if response.status != 200:
                        error = True
                        break
                    times.append((end_time - start_time) * 1000)  # Convert seconds to milliseconds

        return times, error

    async def save_times_results(self, query, times, folder_path):
        os.makedirs(folder_path, exist_ok=True)
        file_path = os.path.join(folder_path, f"{query}.txt")

        async with aiofiles.open(file_path, mode='w') as file:
            for time_value in times:
                await file.write(f"{time_value}\n")

    async def send_queries(self, num_queries, save_file_path):
        with open(self.query_file_path, 'r') as file:
            graphql_queries = json.load(file)

        tasks = []
        for query, graphql_query in graphql_queries.items():
            payload = {"query": graphql_query}
            tasks.append(asyncio.create_task(self.runSingleTestSet(payload, num_queries)))

        results = await asyncio.gather(*tasks)  # Gather results from all tasks

        for query, (times, error) in zip(graphql_queries.keys(), results):
            print("************ ", query, " ************")
            if error:
                print("ERROR")
            else:
                times.sort()
                numpy_times = np.asarray(times, dtype=np.float32)

                mean_val = np.mean(numpy_times).item()
                variance = np.var(numpy_times).item()
                max_val = np.max(numpy_times).item()
                min_val = np.min(numpy_times).item()
                median_val = np.median(numpy_times).item()
                stdev_val = np.std(numpy_times).item()
                percentile_90 = np.percentile(numpy_times, 90).item()

                print("Mean: ", mean_val, "ms")
                # ... (print remaining statistics)

            await self.save_times_results(query, times, save_file_path)

            print()

    async def runStressTest(self, num_queries):
        with open(self.query_file_path, 'r') as file:
            graphql_queries = json.load(file)

        tasks = []
        for query, graphql_query in graphql_queries.items():
            for _ in range(num_queries):
                tasks.append(asyncio.create_task(self.runSingleTestSet({"query": graphql_query}, num_queries)))

        results = await asyncio.gather(*tasks)  # Gather results from all tasks

        return results
class StatisticsTest:
    @staticmethod
    def test_normality(data):
        _, p_value = stats.shapiro(data)
        return p_value
                

    @staticmethod
    def test_exponential(data):
        loc_estimate = min(data)
        scale_estimate = 1 / np.mean(data)
        _, p_value = stats.kstest(data, 'expon', args=(loc_estimate, scale_estimate))
        return p_value


class StatisticsCalculator:
    @staticmethod
    def calculate_statistic(file_path):
        with open(file_path, 'r', encoding='latin-1') as file:
            query_times = [float(line.strip()) for line in file]

        mean_val = np.mean(query_times)
        median_val = np.median(query_times)
        max_val = np.max(query_times)
        min_val = np.min(query_times)
        variance = np.var(query_times)
        std_dev = np.std(query_times)
        percentile_90 = np.percentile(query_times, 90)
        return mean_val, median_val, max_val, min_val, variance, std_dev, percentile_90
    
    def generate_data_dict(folder_paths, statistic='mean'):
        data_dict = {}

        for folder_path in folder_paths:
            folder_data = {}
            for file_name in os.listdir(folder_path):
                if file_name.endswith(".txt"):
                    file_path = os.path.join(folder_path, file_name)
                    stats = StatisticsCalculator.calculate_statistic(file_path)
                    stats_dict = {
                        'mean': stats[0],
                        'median': stats[1],
                        'max': stats[2],
                        'min': stats[3],
                        'variance': stats[4],
                        'standard_deviation': stats[5],
                        'percentile_90': stats[6]
                    }

                    folder_data[os.path.splitext(file_name)[0]] = stats_dict.get(statistic)

            folder_name = os.path.basename(folder_path).replace('queries_times_', '')
            data_dict[folder_name] = folder_data

        return data_dict

class StatisticsGenerator:

    @staticmethod
    def generate_comparison_chart(folder_paths, statistic='mean'):
        data_dict = StatisticsCalculator.generate_data_dict(folder_paths, statistic)

        # Flatten the dictionary to get a list of tuples (folder_name, file_name, stat_value)
        flat_data = [(folder_name, file_name, stat_value) for folder_name, folder_data in data_dict.items() for file_name, stat_value in folder_data.items()]

        # Sort the data alphabetically primarily by file_name and secondarily by folder_name
        sorted_data = sorted(flat_data, key=lambda x: (x[1], x[0]))

        labels = [f"{file_name}" for _, file_name, _ in sorted_data]
        values = [stat_value for _, _, stat_value in sorted_data]
        folder_names = [folder_name for folder_name, _, _ in sorted_data]

        # Sort unique folder names alphabetically for consistency after regenerating
        unique_folder_names = sorted(set(folder_names))

        # Create a color map dictionary based on alphabetical order
        color_map = {folder_name: plt.cm.tab10(i) for i, folder_name in enumerate(unique_folder_names)}
        # Assign colors based on alphabetical order
        colors = [color_map[folder_name] for folder_name in folder_names]

        x = np.arange(len(labels))
        width = 0.3

        fig, ax = plt.subplots(figsize=(10, 5))  # Adjust figure size as needed
        bars = ax.bar(x, values, width, label=statistic.capitalize(), color=colors)

        ax.set_ylabel(f'{statistic.capitalize()} Values')
        ax.set_title(f'Comparison of {statistic.capitalize()}')
        ax.set_xticks(x)
        ax.set_xticklabels(labels, rotation=45, ha='right')

        # Manually create legend handles with appropriate colors
        legend_handles = [plt.Line2D([0], [0], color=color_map[folder_name], lw=4) for folder_name in unique_folder_names]
        legend_labels = [f'{folder}' for folder in unique_folder_names]

        # Add legend with custom handles and labels
        ax.legend(legend_handles, legend_labels)

        # Save the plot to a BytesIO object
        img_buf = BytesIO()
        plt.savefig(img_buf, format='png', dpi=300, bbox_inches='tight', pad_inches=0.3)
        plt.close()

        return img_buf
    
    @staticmethod
    def generate_basic_graph(data, filename):
        # Function implementation
        plt.figure()
        plt.plot(data, marker='o', linestyle='-')
        plt.title(f"Basic Graph - {filename}")
        plt.xlabel("Query")
        plt.ylabel("Time (ms)")
        plt.grid(True)
        # Save the plot to a BytesIO object
        img_buf = BytesIO()
        plt.savefig(img_buf, format='png')
        plt.close()

        # Return the BytesIO object
        return img_buf

    @staticmethod
    def generate_histogram(data, filename):
        plt.figure()
        
        # Square Root Rule
        num_bins = int(np.sqrt(len(data)))

        # Plot the histogram
        plt.hist(data, bins=num_bins, color='blue', edgecolor='black', density=True, alpha=0.7, label='Data Histogram')

        # Overlay normal distribution curve
        mu, sigma = mean(data), stdev(data)
        x = np.linspace(min(data), max(data), 100)
        y = norm.pdf(x, mu, sigma)
        plt.plot(x, y, 'r-', linewidth=2, label='Normal Distribution')

        # Overlay exponential distribution curve
        loc, scale = min(data), stdev(data)  # Adjusted scale parameter
        y_exponential = expon.pdf(x, loc, scale)
        plt.plot(x, y_exponential, 'g-', linewidth=2, label='Exponential Distribution')
        
        plt.title(f"Histogram - {filename}")
        plt.xlabel("Time (ms)")
        plt.ylabel("Frequency")
        plt.grid(True)

        # Save the plot to a BytesIO object
        img_buf = BytesIO()
        plt.savefig(img_buf, format='png')
        plt.close()

        # Return the BytesIO object
        return img_buf

    
    @staticmethod
    def generate_overall_stats(folder_paths):
        
        overall_doc = docx.Document()
        
        # for stat_type in ['mean', 'variance', 'standard_deviation']:
        for stat_type in ['mean', 'median', 'max', 'min', 'variance', 'standard_deviation', 'percentile_90']:

            # Generate comparison chart and add it to the overall doc
            img_buf = StatisticsGenerator.generate_comparison_chart(folder_paths, statistic=stat_type)
            overall_doc.add_heading(f"Comparison Chart - {stat_type.capitalize()}", level=1)
            overall_doc.add_picture(img_buf, width=docx.shared.Inches(7))
            overall_doc.add_page_break()
        
        # overall_output_path = os.path.join("./", "overall_stats.docx")
        # overall_output_path = folder_paths.join("./", "overall_stats.docx")
        overall_output_path = os.path.join(folder_paths[0], "overall_stats.docx")
        overall_doc.save(overall_output_path)

    
    @staticmethod
    def generate_stats_docx(folder_paths):
        for folder_path in folder_paths:
            doc = docx.Document()
            doc.add_heading(f"Statistics for Files in {folder_path}", level=1)

            for filename in os.listdir(folder_path):
                if filename.endswith(".txt"):
                    file_path = os.path.join(folder_path, filename)

                    with open(file_path, 'r') as file:
                        content = file.read().splitlines()

                        # Extract numerical data from the file
                        data = [float(line) for line in content if line.strip().replace('.', '').isdigit()]

                        if data:
                            # Basic statistics
                            doc.add_heading(f"File: {filename}", level=2)
                            doc.add_paragraph(f"Number of queries: {round(len(data), 2)}")
                            doc.add_paragraph(f"Mean: {round(np.mean(data), 2)}")
                            doc.add_paragraph(f"Median: {round(np.median(data), 2)}")
                            doc.add_paragraph(f"Max: {round(np.max(data), 2)}")
                            doc.add_paragraph(f"Min: {round(np.min(data), 2)}")
                            doc.add_paragraph(f"Variance: {round(np.var(data), 2)}")
                            doc.add_paragraph(f"Standard Deviation: {round(np.std(data), 2)}")
                            doc.add_paragraph(f"90th Percentile: {round(np.percentile(data, 90), 2)}")
                            
                            # Perform normality test
                            p_value_normality = StatisticsTest.test_normality(data)
                            doc.add_paragraph(f"Is normal: {'Yes' if p_value_normality > 0.05 else 'No'}")
                            doc.add_paragraph(f"\tShapiro-Wilk test p-value: {p_value_normality}")

                            # Perform exponential distribution test
                            p_value_exponential = StatisticsTest.test_exponential(data)
                            doc.add_paragraph(f"Is exponential: {'Yes' if p_value_exponential > 0.05 else 'No'}")
                            doc.add_paragraph(f"\tKolmogorov-Smirnov test p-value: {p_value_exponential}")
                            doc.add_paragraph("\n")

                            #print out the paragraph
                            

                            # Generate basic graph
                            img_buf_basic = StatisticsGenerator.generate_basic_graph(data, filename)
                            doc.add_picture(img_buf_basic, width=docx.shared.Inches(6))

                            # Generate histogram
                            img_buf_histogram = StatisticsGenerator.generate_histogram(data, filename)
                            doc.add_picture(img_buf_histogram, width=docx.shared.Inches(6))
                            doc.add_page_break()

            output_docx_path = os.path.join(folder_path, "statistics.docx")
            doc.save(output_docx_path)




query_file_path = 'D:/Documents/Unob_7/STC/STC_code/DB_Performance/_Test/Performance/Queries/gql_queries.json'
save_file_path = 'D:/Documents/Unob_7/STC/STC_code/DB_Performance/_Test/Performance/Test_v1/postgres'

async def main():
    
    query_performance_analyzer = QueryPerformanceAnalyzer(query_file_path)
    await query_performance_analyzer.send_queries(5, save_file_path)

loop = asyncio.get_event_loop()
loop.run_until_complete(main())
loop.close()


StatisticsGenerator.generate_overall_stats([save_file_path])
StatisticsGenerator.generate_stats_docx([save_file_path])